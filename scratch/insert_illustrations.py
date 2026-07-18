import docx
from docx.shared import Inches
import os

doc_path = r"C:\Users\Kien\Downloads\Điện thoại thần kỳ.docx"
output_path = r"C:\Users\Kien\Downloads\Điện thoại thần kỳ (đã chèn ảnh).docx"

# Image paths
images = {
    "ch17_kiss": r"C:\Users\Kien\.gemini\antigravity\brain\06f9f4b1-b4d6-4e77-8027-66b2b117cd8a\ch17_kiss_1784315253576.png",
    "ch17_crawl": r"C:\Users\Kien\.gemini\antigravity\brain\06f9f4b1-b4d6-4e77-8027-66b2b117cd8a\ch17_crawl_1784315263084.png",
    "ch17_playful": r"C:\Users\Kien\.gemini\antigravity\brain\06f9f4b1-b4d6-4e77-8027-66b2b117cd8a\ch17_playful_1784315272783.png",
    "ch18_moonlight": r"C:\Users\Kien\.gemini\antigravity\brain\06f9f4b1-b4d6-4e77-8027-66b2b117cd8a\ch18_moonlight_1784315282187.png",
    "ch18_jade": r"C:\Users\Kien\.gemini\antigravity\brain\06f9f4b1-b4d6-4e77-8027-66b2b117cd8a\ch18_jade_1784315293299.png",
    "ch18_cozy": r"C:\Users\Kien\.gemini\antigravity\brain\06f9f4b1-b4d6-4e77-8027-66b2b117cd8a\ch18_cozy_1784315302455.png",
    "ch19_letter": r"C:\Users\Kien\.gemini\antigravity\brain\06f9f4b1-b4d6-4e77-8027-66b2b117cd8a\ch19_letter_1784315311858.png",
    "ch19_spear": r"C:\Users\Kien\.gemini\antigravity\brain\06f9f4b1-b4d6-4e77-8027-66b2b117cd8a\ch19_spear_1784315321170.png",
    "ch19_soccer": r"C:\Users\Kien\.gemini\antigravity\brain\06f9f4b1-b4d6-4e77-8027-66b2b117cd8a\ch19_soccer_1784315329769.png"
}

doc = docx.Document(doc_path)

def add_image_after_paragraph(p, img_path):
    p_img = p.insert_paragraph_before()
    p_img.alignment = 1 # Center alignment
    r = p_img.add_run()
    r.add_picture(img_path, width=Inches(4.5))
    # Add an empty spacing paragraph
    p_spacing = p.insert_paragraph_before()
    p_spacing.add_run().text = ""

# Track the context we want to search for
found_counts = {k: 0 for k in images.keys()}

# We scan paragraphs and find matches to insert images before/after them
for i in range(len(doc.paragraphs)):
    text = doc.paragraphs[i].text.strip()
    
    # Chapter 17 Matchers
    if "một cái hôn sâu của đôi tình nhân..." in text and found_counts["ch17_kiss"] == 0:
        add_image_after_paragraph(doc.paragraphs[i], images["ch17_kiss"])
        found_counts["ch17_kiss"] = 1
        
    elif "rón rén trườn xuống khỏi phiến đá" in text and found_counts["ch17_crawl"] == 0:
        add_image_after_paragraph(doc.paragraphs[i], images["ch17_crawl"])
        found_counts["ch17_crawl"] = 1
        
    elif "Chắc là không dám đâu!!! Tiểu hồ ly" in text and found_counts["ch17_playful"] == 0:
        add_image_after_paragraph(doc.paragraphs[i], images["ch17_playful"])
        found_counts["ch17_playful"] = 1
        
    # Chapter 18 Matchers
    elif "Khung cảnh quá mực tuyệt diệu trong đêm" in text and found_counts["ch18_moonlight"] == 0:
        add_image_after_paragraph(doc.paragraphs[i], images["ch18_moonlight"])
        found_counts["ch18_moonlight"] = 1
        
    elif "quà của phu quân" in text and found_counts["ch18_jade"] == 0:
        add_image_after_paragraph(doc.paragraphs[i], images["ch18_jade"])
        found_counts["ch18_jade"] = 1
        
    elif "Kiệt tham lam còn ụp mặt vào ngực" in text and found_counts["ch18_cozy"] == 0:
        add_image_after_paragraph(doc.paragraphs[i], images["ch18_cozy"])
        found_counts["ch18_cozy"] = 1
        
    # Chapter 19 Matchers
    elif "Đông Nam Á Học Viện sao? Được, Huyên" in text and found_counts["ch19_letter"] == 0:
        add_image_after_paragraph(doc.paragraphs[i], images["ch19_letter"])
        found_counts["ch19_letter"] = 1
        
    elif "đầu thương sắc lẹm, tỏa ra sát khí nhiếp người" in text and found_counts["ch19_spear"] == 0:
        add_image_after_paragraph(doc.paragraphs[i], images["ch19_spear"])
        found_counts["ch19_spear"] = 1
        
    elif "Euro, giải đấu của những bậc thầy đá bóng" in text and found_counts["ch19_soccer"] == 0:
        add_image_after_paragraph(doc.paragraphs[i], images["ch19_soccer"])
        found_counts["ch19_soccer"] = 1

doc.save(output_path)
print("Finished. Image insertion report:")
print(found_counts)
