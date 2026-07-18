import docx
from docx.shared import Inches
import os

doc_path = r"C:\Users\Kien\Downloads\Điện thoại thần kỳ.docx"
output_path = r"C:\Users\Kien\Downloads\Điện thoại thần kỳ (đã chèn ảnh ch29).docx"

# Seductive prompts definition for Chapter 29 Lập Hạ
prompts = {
    "ch29_lapha_squeeze": (
        "Seductive and extremely sexy anime style illustration, ancient Vietnamese/Chinese cultivation. "
        "A gorgeous, highly alluring young female cultivator named Lap Ha with long hair, wearing a very "
        "revealing, cleavage-baring light blue traditional robe that showcases her voluptuous figure, huge breasts, "
        "and slender waist. She is gracefully squeezing through a narrow rocky cave gap, looking extremely attractive. "
        "Mystical purple glowing mist in the background. High detailed fantasy art."
    ),
    "ch29_lapha_blush": (
        "Seductive and extremely sexy anime style illustration, ancient cultivation setting. A beautiful, curvaceous "
        "young female cultivator named Lap Ha with long hair, wearing a revealing, tight-fitting ancient robe highlighting "
        "her deep cleavage. Her face is blushing red with embarrassment as a handsome young cultivator stares at her intensely. "
        "Dangerous purple-glowing dragon silhouette in the misty background. Seductive romantic tension."
    ),
    "ch29_lapha_charm": (
        "Seductive and extremely sexy anime style illustration, ancient cultivation setting. A voluptuous, gorgeous "
        "young female cultivator named Lap Ha with an hourglass figure, wearing a revealing, cleavage-baring robe. "
        "She looks slightly annoyed yet extremely seductive, crossing her arms. A slick guy with a playful smirk is talking to her, "
        "trying to charm her. Detailed environment, fantasy atmosphere."
    )
}

# Image paths (Assuming they will be generated or are present)
images = {
    "ch29_lapha_squeeze": r"C:\Users\Kien\.gemini\antigravity\brain\06f9f4b1-b4d6-4e77-8027-66b2b117cd8a\ch29_lapha_squeeze.png",
    "ch29_lapha_blush": r"C:\Users\Kien\.gemini\antigravity\brain\06f9f4b1-b4d6-4e77-8027-66b2b117cd8a\ch29_lapha_blush.png",
    "ch29_lapha_charm": r"C:\Users\Kien\.gemini\antigravity\brain\06f9f4b1-b4d6-4e77-8027-66b2b117cd8a\ch29_lapha_charm.png"
}

print("Seductive prompts for Lap Ha in Chapter 29:")
for k, v in prompts.items():
    print(f"\n[{k}]: {v}")

if os.path.exists(doc_path):
    doc = docx.Document(doc_path)
    found_counts = {k: 0 for k in images.keys()}
    
    def add_image_after_paragraph(p, img_path):
        if not os.path.exists(img_path):
            print(f"Warning: Image file not found: {img_path}. Adding placeholder text instead.")
            p_img = p.insert_paragraph_before()
            p_img.alignment = 1
            r = p_img.add_run()
            r.text = f"[ẢNH MINH HỌA SEXY LẬP HẠ - CHƯA CÓ FILE: {os.path.basename(img_path)}]"
            r.bold = True
            return
            
        p_img = p.insert_paragraph_before()
        p_img.alignment = 1 # Center
        r = p_img.add_run()
        r.add_picture(img_path, width=Inches(4.5))
        p_spacing = p.insert_paragraph_before()
        p_spacing.add_run().text = ""

    for i in range(len(doc.paragraphs)):
        text = doc.paragraphs[i].text.strip()
        
        # Matchers for Chapter 29
        if "tiến thẳng vào trung tâm." in text and found_counts["ch29_lapha_squeeze"] == 0:
            add_image_after_paragraph(doc.paragraphs[i], images["ch29_lapha_squeeze"])
            found_counts["ch29_lapha_squeeze"] = 1
            
        elif "Lập Hạ đỏ mặt, gắt lên" in text and found_counts["ch29_lapha_blush"] == 0:
            add_image_after_paragraph(doc.paragraphs[i], images["ch29_lapha_blush"])
            found_counts["ch29_lapha_blush"] = 1
            
        elif "Đi chung với anh, anh bảo kê từ A đến Z" in text and found_counts["ch29_lapha_charm"] == 0:
            add_image_after_paragraph(doc.paragraphs[i], images["ch29_lapha_charm"])
            found_counts["ch29_lapha_charm"] = 1

    doc.save(output_path)
    print("\nSaved document with placeholders successfully.")
    print("Report:", found_counts)
else:
    print("Error: Source document not found at", doc_path)
